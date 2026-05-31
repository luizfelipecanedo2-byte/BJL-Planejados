from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Configuração de estilo padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Cabeçalho
    header1 = doc.add_paragraph()
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header1.add_run('UENF - Universidade Estadual do Norte Fluminense Darcy Ribeiro\n')
    run.bold = True
    run.font.size = Pt(12)
    
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header2.add_run('LEEL - Laboratório de Estudos da Educação e Linguagem\n')
    run.bold = True
    header2.add_run('Núcleo de Estágio da UENF')

    doc.add_paragraph() # Espaço

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('FICHA DE OBSERVAÇÃO DE ATIVIDADE e APRECIAÇÃO DE DOCÊNCIA')
    run.bold = True
    run.font.size = Pt(12)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('(Para o estagiário observar 2 atividades e 2 aulas diferentes)')

    doc.add_paragraph() # Espaço

    # Informações Gerais
    info_fields = [
        ('Instituição:', 'Escola Jardim de Infância Ana Passos'),
        ('Disciplina:', 'Português'),
        ('Série:', 'Pré II'),
        ('Professor:', 'Larissa Gonçalves dos Santos Gomes'),
        ('Data da observação:', '16/04/2026'),
        ('Duração da aula:', '4 horas')
    ]

    for label, value in info_fields:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f' {value}')

    doc.add_paragraph() # Espaço

    # Planejamento
    p = doc.add_paragraph()
    p.add_run('Planejamento:').bold = True
    
    p = doc.add_paragraph('Percebe-se a presença de objetivos previamente determinados, de acordo com os campos de Experiência (BNCC)? Quais?')
    p = doc.add_paragraph()
    p.add_run('( x ) sim ( ) não').bold = True
    
    p = doc.add_paragraph()
    p.add_run('( x ) Escuta, fala, pensamento e imaginação; ( x ) O eu, o outro e o nós ( ) Corpo, gestos e movimento; ( ) Traços, sons, cores e formas; ( ) Espaços, tempos, quantidades, relações e transformações')

    p = doc.add_paragraph()
    p.add_run('Nota-se a exigência de um planejamento?')
    p = doc.add_paragraph()
    p.add_run('( x ) sim ( ) não').bold = True

    doc.add_paragraph()

    # Conteúdo
    p = doc.add_paragraph()
    p.add_run('Conteúdo:').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Assunto: ').bold = True
    p.add_run('Povos indígenas.')
    
    p = doc.add_paragraph()
    p.add_run('Tópicos principais: ').bold = True
    p.add_run('Cultura e costumes dos povos indígenas.')

    doc.add_paragraph()

    # Estratégias
    p = doc.add_paragraph()
    p.add_run('Estratégias:').bold = True
    
    p = doc.add_paragraph('1. O assunto foi introduzido através de:')
    p = doc.add_paragraph('( ) exposição pelo professor\n( ) leitura de texto\n( x ) perguntas dirigidas à turma\n( ) Outras atividades. Quais? _________________________________________________')

    p = doc.add_paragraph('2. Procedimentos pedagógicos empregados no desenvolvimento do assunto:')
    p = doc.add_paragraph('( ) elaboração de atividades em conjunto com uma turma\n( x ) exposição pelo professor\n( ) exposição e debate simultâneo com a turma\n( ) técnicas de dinâmica de grupo\n( ) atividades dos alunos sob a supervisão do professor\n( ) demonstração\n( ) outras. Quais? _________________________________________________')

    doc.add_paragraph()

    # Recursos
    p = doc.add_paragraph('3. Recursos didáticos utilizados:')
    p = doc.add_paragraph('( ) quadro de giz\n( ) reália, modelos\n( x ) material impresso\n( ) cartazes\n( ) gravações\n( ) ilustrações\n( ) álbum seriado projeções de: ( ) slides ( ) filmes ( ) transparências\n( ) outras. Quais? _________________________________________________')

    doc.add_paragraph()

    # Atividades
    p = doc.add_paragraph()
    p.add_run('Atividades desenvolvidas durante a observação. O envolvimento da turma durante a atividade foi:').bold = True
    p = doc.add_paragraph('( x ) excelente\n( ) muito bom\n( ) bom\n( ) regular\n( ) insuficiente')

    doc.add_paragraph()

    # Professor
    p = doc.add_paragraph()
    p.add_run('Professor:').bold = True
    
    prof_questions = [
        'Apresenta o conteúdo com dinamicidade?',
        'Mantinha bom relacionamento com a turma?',
        'Apresentava domínio do conteúdo?',
        'Apresentava explicações claras?',
        'Solicitava a participação dos alunos?',
        'Tornava a explicar quando solicitado?'
    ]
    
    for q in prof_questions:
        p = doc.add_paragraph(q)
        p = doc.add_paragraph('( x ) sim ( ) não ( ) em parte')

    doc.add_paragraph()

    # Avaliação
    p = doc.add_paragraph()
    p.add_run('Avaliação:').bold = True
    
    p = doc.add_paragraph('Houve preocupação por parte do professor em avaliar a atividade proposta?')
    p = doc.add_paragraph('( x ) sim ( ) não')
    
    p = doc.add_paragraph('Se você tomou conhecimento dos objetivos, havia relação entre a avaliação e o que foi observado?')
    p = doc.add_paragraph('( x ) sim ( ) não ( ) em parte')

    p = doc.add_paragraph('No caso de ter havido avaliação, foram empregados os seguintes Instrumentos:')
    p = doc.add_paragraph('( x ) interrogatório\n( ) teste escrito\n( ) debate\n( ) elaboração de tarefas, exercícios ou trabalhos práticos.\n( ) outros. Quais? _________________________________________________')

    doc.add_paragraph()

    # Comentários
    p = doc.add_paragraph()
    p.add_run('Registre o(s) aspecto(s) da aula que mais chamou(ram) sua atenção:').bold = True
    p = doc.add_paragraph('A aula foi introduzida de maneira dinâmica, divertida e que prendeu a atenção das crianças e despertou o interesse e a curiosidade. Foi dando a eles pipoca e explicado que esse costume que adquirimos no Brasil foi os povos indígenas que nos ensinaram como muitos outros alimentos do nosso dia a dia.')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Se algum aluno despertou sua atenção de maneira especial, registre o fato.').bold = True
    p = doc.add_paragraph('A turma no geral chamou a atenção pela participação e interação na aula.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Assinatura
    p = doc.add_paragraph()
    p.add_run('_________________________________________________\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Assinatura do estagiário: Felipe Lourenço Fonseca Canêdo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('Ficha_Observacao_UENF.docx')
    print("Arquivo 'Ficha_Observacao_UENF.docx' criado com sucesso!")

if __name__ == "__main__":
    create_document()
